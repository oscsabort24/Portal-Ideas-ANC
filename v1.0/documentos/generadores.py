"""Generadores de los 6 documentos formales en .docx (python-docx).

Cada función recibe el dict de `datos` ya ensamblado por
documentos/service.py (campos estructurales de la BD + campos
narrativos del stub de IA) y escribe el .docx en `ruta_salida`.

Todo campo usa `.get(clave) or "valor por defecto"` — nunca se escribe
una celda/párrafo vacío en el documento, mismo patrón de fallback
textual que usaba el prototipo v0.1 ("Pendiente de definir", "Por
definir", etc.).

BPMN as-is/to-be se genera como diagrama visual real (cajas y flechas)
con Graphviz — ver _generar_diagrama_bpmn. Requiere el ejecutable
`dot` de Graphviz instalado en el sistema (no solo el paquete de
Python); ver _asegurar_graphviz_disponible para el manejo de ese caso.

Estilo visual: se replica lo más posible la misma paleta que las
plantillas HTML (documentos/plantillas_html.py) usando shading de
celda y color de texto vía XML de python-docx (docx.oxml) — el
formato .docx no soporta CSS, así que esto NO es 1:1 con el HTML/PDF,
pero comparte colores y estructura para que las tres versiones
(preview, PDF, Word) se sientan parte del mismo sistema.
"""

import os
import platform
import shutil

import graphviz
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

_AZUL = RGBColor(0x2E, 0x5F, 0xAA)
_AZUL_CLARO_HEX = "EAF1FB"
_AZUL_HEX = "2E5FAA"
_NARANJA = RGBColor(0xE8, 0x76, 0x2C)
_VERDE = RGBColor(0x00, 0x71, 0x3D)
_VERDE_CLARO_HEX = "E8F5ED"
_NARANJA_OSCURO = RGBColor(0xB4, 0x53, 0x09)
_NARANJA_CLARO_HEX = "FEF3C7"
_BLANCO = RGBColor(0xFF, 0xFF, 0xFF)

_RACI_COLORES = {
    "R": ("E8F4FD", RGBColor(0x2E, 0x5F, 0xAA)),
    "A": (_VERDE_CLARO_HEX, _VERDE),
    "C": (_NARANJA_CLARO_HEX, _NARANJA_OSCURO),
    "I": ("F0F0EF", RGBColor(0x6B, 0x72, 0x80)),
}


def _sombrear_celda(cell, color_hex: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def _colorear_texto_celda(cell, color: RGBColor, negrita: bool = True, centrado: bool = False) -> None:
    for parrafo in cell.paragraphs:
        if centrado:
            parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in parrafo.runs:
            run.font.color.rgb = color
            run.font.bold = negrita


def _borde_inferior(paragraph, color_hex: str, grosor: int = 18) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    borde = OxmlElement("w:bottom")
    borde.set(qn("w:val"), "single")
    borde.set(qn("w:sz"), str(grosor))
    borde.set(qn("w:space"), "4")
    borde.set(qn("w:color"), color_hex)
    pBdr.append(borde)
    pPr.append(pBdr)


def _agregar_encabezado(doc: Document, titulo: str, subtitulo: str) -> None:
    p_titulo = doc.add_heading(titulo, level=0)
    for run in p_titulo.runs:
        run.font.color.rgb = _AZUL
    _borde_inferior(p_titulo, _AZUL_HEX, grosor=24)

    p = doc.add_paragraph(subtitulo)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = _NARANJA
        run.font.bold = True
    doc.add_paragraph()


def _seccion(doc: Document, titulo: str, contenido: str) -> None:
    p_titulo = doc.add_heading(titulo, level=1)
    for run in p_titulo.runs:
        run.font.color.rgb = _AZUL
    _borde_inferior(p_titulo, _AZUL_CLARO_HEX)
    doc.add_paragraph(contenido or "Pendiente de definir")


def generar_charter_docx(datos: dict, ruta_salida: str) -> None:
    """datos: nombre_proyecto, area_solicitante, solicitante, fecha_emision,
    programa, procedimiento_sig, justificacion_alcance, objetivos,
    beneficios_esperados, principales_entregables,
    riesgos_identificados: list[{riesgo, mitigacion}], estado."""
    doc = Document()
    _agregar_encabezado(doc, datos.get("nombre_proyecto") or "Project Charter", datos.get("programa") or "Transformación Digital")

    tabla_meta = doc.add_table(rows=0, cols=2)
    tabla_meta.style = "Table Grid"
    for etiqueta, clave, default in [
        ("Área Solicitante", "area_solicitante", "No especificada"),
        ("Solicitante", "solicitante", "Participante"),
        ("Fecha de Emisión", "fecha_emision", "—"),
        ("Procedimiento SIG", "procedimiento_sig", "No existe"),
    ]:
        fila = tabla_meta.add_row().cells
        fila[0].text = etiqueta
        fila[1].text = datos.get(clave) or default
        _sombrear_celda(fila[0], _AZUL_CLARO_HEX)
        _colorear_texto_celda(fila[0], _AZUL)
    doc.add_paragraph()

    _seccion(doc, "Justificación y Alcance del Proyecto", datos.get("justificacion_alcance"))
    _seccion(doc, "Objetivos del Proyecto", datos.get("objetivos"))
    _seccion(doc, "Beneficios Esperados", datos.get("beneficios_esperados"))
    _seccion(doc, "Principales Entregables", datos.get("principales_entregables"))

    p_titulo_riesgos = doc.add_heading("Riesgos Identificados", level=1)
    for run in p_titulo_riesgos.runs:
        run.font.color.rgb = _AZUL
    _borde_inferior(p_titulo_riesgos, _AZUL_CLARO_HEX)

    riesgos = datos.get("riesgos_identificados") or []
    if riesgos:
        tabla_riesgos = doc.add_table(rows=1, cols=2)
        tabla_riesgos.style = "Table Grid"
        encabezado = tabla_riesgos.rows[0].cells
        encabezado[0].text = "Riesgo identificado"
        encabezado[1].text = "Estrategia de mitigación"
        for celda in encabezado:
            _sombrear_celda(celda, _AZUL_HEX)
            _colorear_texto_celda(celda, _BLANCO)
        for i, r in enumerate(riesgos):
            fila = tabla_riesgos.add_row().cells
            fila[0].text = r.get("riesgo") or "Pendiente de definir"
            fila[1].text = r.get("mitigacion") or "Pendiente de definir"
            if i % 2 == 1:
                for celda in fila:
                    _sombrear_celda(celda, _AZUL_CLARO_HEX)
    else:
        doc.add_paragraph("Pendiente de definir")

    doc.add_paragraph()
    p_estado = doc.add_paragraph()
    p_estado.add_run("Estado del documento: ").bold = True
    estado_texto = datos.get("estado") or "—"
    run_estado = p_estado.add_run(estado_texto)
    run_estado.bold = True
    run_estado.font.color.rgb = _VERDE if "listo" in estado_texto.lower() else _NARANJA_OSCURO

    doc.save(ruta_salida)


def _asegurar_graphviz_disponible() -> None:
    """Verifica que el ejecutable `dot` de Graphviz esté disponible.

    1. Si ya está en el PATH (cualquier SO), no hace nada.
    2. Si no, intenta agregar rutas candidatas conocidas SEGÚN EL SISTEMA
       OPERATIVO — hoy solo Windows, porque en Linux/macOS Graphviz
       instalado vía el gestor de paquetes (apt/yum/brew) ya queda en el
       PATH del sistema sin necesitar este workaround.
    3. Si sigue sin encontrarse, lanza un RuntimeError descriptivo (en vez
       del ExecutableNotFound genérico de la librería graphviz) que dice
       exactamente qué falta y cómo resolverlo.
    """
    if shutil.which("dot"):
        return

    rutas_candidatas = {
        "Windows": [r"C:\Program Files\Graphviz\bin"],
    }.get(platform.system(), [])

    for ruta in rutas_candidatas:
        if os.path.isdir(ruta):
            os.environ["PATH"] += os.pathsep + ruta
            if shutil.which("dot"):
                return

    raise RuntimeError(
        "Graphviz (el ejecutable 'dot') no está instalado o no está en el PATH "
        "de este proceso. Instálalo desde https://graphviz.org/download/ "
        "(Windows) o con el gestor de paquetes de tu sistema (ej. "
        "'apt install graphviz' en Linux), y asegúrate de que el directorio "
        "bin quede en el PATH — luego reinicia el servidor para que tome el "
        "PATH actualizado."
    )


_FORMA_POR_TIPO = {
    "inicio": "ellipse",
    "fin": "ellipse",
    "tarea": "box",
    "decision": "diamond",
}


def _generar_diagrama_bpmn(pasos: list[dict], ruta_base_sin_extension: str) -> str | None:
    """Genera un PNG (cajas y flechas) para un proceso as-is o to-be.

    El orden de las cajas es el orden del array `pasos` (por índice) —
    NO el campo "id": el stub de IA hoy ni siquiera lo incluye, y no hay
    garantía de que sea secuencial cuando exista en datos reales.

    Devuelve la ruta del PNG generado, o None si `pasos` viene vacío —
    en ese caso el llamador debe usar el fallback de texto "Pendiente de
    definir", igual que el resto de los campos del documento. Un solo
    paso también es válido: se dibuja un único nodo sin flechas (es
    justamente lo que produce el stub de IA hoy).
    """
    if not pasos:
        return None

    _asegurar_graphviz_disponible()

    grafo = graphviz.Digraph(format="png")
    # rankdir="TB" (vertical): con "LR" el diagrama crecía casi solo en
    # ancho (una tira horizontal aplastada e ilegible al insertarse a un
    # ancho fijo de 6.5in en el .docx) — confirmado generando el PNG real
    # con ambos valores antes de aplicar este cambio.
    grafo.attr(rankdir="TB")
    grafo.attr("node", fontname="Helvetica", fontsize="11")
    grafo.attr("edge", color=f"#{_AZUL_HEX}")

    for i, paso in enumerate(pasos):
        actor = paso.get("actor") or "—"
        accion = paso.get("accion") or "Pendiente de definir"
        forma = _FORMA_POR_TIPO.get(paso.get("tipo"), "box")
        if paso.get("usado_en_to_be", True):
            estilo = {"style": "filled", "fillcolor": f"#{_AZUL_CLARO_HEX}", "color": f"#{_AZUL_HEX}", "fontcolor": "#22282E"}
        else:
            # Paso del AS-IS que el rediseño TO-BE elimina: gris apagado y
            # borde punteado para distinguirlo de un paso normal.
            estilo = {"style": "filled,dashed", "fillcolor": "#F0F0EF", "color": "#6B7280", "fontcolor": "#6B7280"}
        grafo.node(f"n{i}", f"{actor}: {accion}", shape=forma, **estilo)

    for i in range(len(pasos) - 1):
        grafo.edge(f"n{i}", f"n{i + 1}")

    return grafo.render(ruta_base_sin_extension, cleanup=True)


def generar_bpmn_docx(datos: dict, ruta_salida: str) -> None:
    """datos: titulo, descripcion, actores: list[str],
    pasos_as_is / pasos_to_be: list[{actor, accion, tipo}].

    Genera un diagrama real por lado (as-is/to-be) con Graphviz y lo
    inserta como imagen — ver _generar_diagrama_bpmn."""
    doc = Document()
    _agregar_encabezado(doc, datos.get("titulo") or "Diagrama de Proceso", "BPMN")

    doc.add_paragraph(datos.get("descripcion") or "Pendiente de definir")

    doc.add_heading("Participantes del proceso", level=1)
    actores = datos.get("actores") or []
    doc.add_paragraph(", ".join(actores) if actores else "No especificados")

    base_sin_extension = os.path.splitext(ruta_salida)[0]

    def _seccion_proceso(titulo_seccion: str, pasos: list[dict], sufijo: str) -> None:
        doc.add_heading(titulo_seccion, level=1)
        ruta_png = _generar_diagrama_bpmn(pasos, f"{base_sin_extension}_{sufijo}")
        if ruta_png:
            doc.add_picture(ruta_png, width=Inches(6.5))
            os.remove(ruta_png)  # ya insertado en el .docx, no hace falta conservarlo en disco
        else:
            doc.add_paragraph("Pendiente de definir")

    _seccion_proceso("Proceso Actual (AS-IS)", datos.get("pasos_as_is") or [], "as_is")
    doc.add_paragraph()
    _seccion_proceso("Proceso Futuro (TO-BE)", datos.get("pasos_to_be") or [], "to_be")

    doc.save(ruta_salida)


def generar_onepager_docx(datos: dict, ruta_salida: str) -> None:
    """datos: titulo, problema, solucion, beneficios: list[str],
    impacto, esfuerzo, proximo_paso."""
    doc = Document()
    _agregar_encabezado(doc, datos.get("titulo") or "One-Pager", "One-pager")

    _seccion(doc, "El problema", datos.get("problema"))
    _seccion(doc, "La solución", datos.get("solucion"))

    p_titulo_beneficios = doc.add_heading("Beneficios esperados", level=1)
    for run in p_titulo_beneficios.runs:
        run.font.color.rgb = _AZUL
    _borde_inferior(p_titulo_beneficios, _AZUL_CLARO_HEX)

    beneficios = datos.get("beneficios") or []
    if beneficios:
        for b in beneficios:
            p_beneficio = doc.add_paragraph(style="List Bullet")
            run = p_beneficio.add_run(b)
            run.font.color.rgb = _VERDE
            run.font.bold = True
    else:
        doc.add_paragraph("No especificados")

    tabla_impacto = doc.add_table(rows=0, cols=2)
    tabla_impacto.style = "Table Grid"
    for etiqueta, clave in [("Impacto", "impacto"), ("Esfuerzo", "esfuerzo")]:
        fila = tabla_impacto.add_row().cells
        fila[0].text = etiqueta
        fila[1].text = datos.get(clave) or "Por definir"
        _sombrear_celda(fila[0], _AZUL_CLARO_HEX)
        _colorear_texto_celda(fila[0], _AZUL)
    doc.add_paragraph()

    _seccion(doc, "Próximo paso recomendado", datos.get("proximo_paso"))

    doc.save(ruta_salida)


def generar_raci_docx(datos: dict, ruta_salida: str) -> None:
    """datos: titulo, actividades: list[{actividad, roles: {rol: R|A|C|I}}],
    leyenda: dict[str, str]."""
    doc = Document()
    _agregar_encabezado(doc, datos.get("titulo") or "Matriz RACI", "RACI")

    actividades = datos.get("actividades") or []
    roles = sorted({rol for a in actividades for rol in (a.get("roles") or {}).keys()})

    tabla = doc.add_table(rows=1, cols=1 + len(roles))
    tabla.style = "Table Grid"
    encabezado = tabla.rows[0].cells
    encabezado[0].text = "Actividad"
    for i, rol in enumerate(roles, start=1):
        encabezado[i].text = rol
    for celda in encabezado:
        _sombrear_celda(celda, _AZUL_HEX)
        _colorear_texto_celda(celda, _BLANCO, centrado=True)

    for a in actividades:
        fila = tabla.add_row().cells
        fila[0].text = a.get("actividad") or "Pendiente de definir"
        valores_rol = a.get("roles") or {}
        for i, rol in enumerate(roles, start=1):
            valor = valores_rol.get(rol)
            fila[i].text = valor or "—"
            if valor and valor in _RACI_COLORES:
                color_hex, color_rgb = _RACI_COLORES[valor]
                _sombrear_celda(fila[i], color_hex)
                _colorear_texto_celda(fila[i], color_rgb, centrado=True)

    doc.add_paragraph()
    p_titulo_leyenda = doc.add_heading("Leyenda", level=1)
    for run in p_titulo_leyenda.runs:
        run.font.color.rgb = _AZUL
    _borde_inferior(p_titulo_leyenda, _AZUL_CLARO_HEX)

    leyenda = datos.get("leyenda") or {}
    for clave, descripcion in leyenda.items():
        p_leyenda = doc.add_paragraph(style="List Bullet")
        run_clave = p_leyenda.add_run(f"{clave}: ")
        run_clave.bold = True
        if clave in _RACI_COLORES:
            run_clave.font.color.rgb = _RACI_COLORES[clave][1]
        p_leyenda.add_run(descripcion)

    doc.save(ruta_salida)


def _celda_bmc(cell, *bloques: tuple[str, str | None]) -> None:
    """Escribe 1 o 2 bloques (título en azul + contenido) dentro de la
    MISMA celda de tabla, apilados uno debajo del otro — así se logra el
    layout "2 bloques en la misma columna" (Actividades+Recursos,
    Relaciones+Canales) sin necesitar sub-tablas anidadas."""
    primer_parrafo = cell.paragraphs[0]
    for indice, (etiqueta, texto) in enumerate(bloques):
        parrafo_titulo = primer_parrafo if indice == 0 else cell.add_paragraph()
        run_titulo = parrafo_titulo.add_run(etiqueta.upper())
        run_titulo.bold = True
        run_titulo.font.color.rgb = _AZUL
        run_titulo.font.size = Pt(9)

        parrafo_contenido = cell.add_paragraph()
        parrafo_contenido.add_run(texto or "Pendiente de definir")

        if indice < len(bloques) - 1:
            cell.add_paragraph()  # separador visual entre los 2 bloques apilados


def generar_bmc_docx(datos: dict, ruta_salida: str) -> None:
    """datos: titulo + los 9 bloques del Business Model Canvas (texto libre
    c/u). Layout en 2 tablas (5 columnas arriba, 2 columnas anchas abajo)
    replicando el grid del canvas — Word no soporta CSS Grid, así que
    "Actividades+Recursos" y "Relaciones+Canales" se apilan como párrafos
    dentro de la misma celda (ver _celda_bmc) en vez de ser 2 filas reales."""
    doc = Document()
    _agregar_encabezado(doc, datos.get("titulo") or "Business Model Canvas", "BMC")

    tabla_arriba = doc.add_table(rows=1, cols=5)
    tabla_arriba.style = "Table Grid"
    celdas = tabla_arriba.rows[0].cells
    _celda_bmc(celdas[0], ("Socios Clave", datos.get("socios_clave")))
    _celda_bmc(celdas[1], ("Actividades Clave", datos.get("actividades_clave")), ("Recursos Clave", datos.get("recursos_clave")))
    _celda_bmc(celdas[2], ("Propuesta de Valor", datos.get("propuesta_valor")))
    _celda_bmc(celdas[3], ("Relaciones con Clientes", datos.get("relaciones_clientes")), ("Canales", datos.get("canales")))
    _celda_bmc(celdas[4], ("Segmentos de Clientes", datos.get("segmentos_clientes")))

    doc.add_paragraph()

    tabla_abajo = doc.add_table(rows=1, cols=2)
    tabla_abajo.style = "Table Grid"
    celdas_abajo = tabla_abajo.rows[0].cells
    _celda_bmc(celdas_abajo[0], ("Estructura de Costos", datos.get("estructura_costos")))
    _celda_bmc(celdas_abajo[1], ("Fuentes de Ingreso", datos.get("fuentes_ingreso")))

    doc.save(ruta_salida)


def generar_business_case_docx(datos: dict, ruta_salida: str) -> None:
    """datos: titulo, resumen_ejecutivo, problema, solucion_propuesta,
    costo_estimado, beneficio_estimado, roi_estimado, payback_estimado,
    supuestos: list[str], recomendacion: Go|No Go|Pendiente de análisis."""
    doc = Document()
    _agregar_encabezado(doc, datos.get("titulo") or "Business Case", "Business Case")

    _seccion(doc, "Resumen Ejecutivo", datos.get("resumen_ejecutivo"))
    _seccion(doc, "El problema", datos.get("problema"))
    _seccion(doc, "Solución propuesta", datos.get("solucion_propuesta"))

    p_titulo_metricas = doc.add_heading("Métricas financieras", level=1)
    for run in p_titulo_metricas.runs:
        run.font.color.rgb = _AZUL
    _borde_inferior(p_titulo_metricas, _AZUL_CLARO_HEX)

    tabla = doc.add_table(rows=0, cols=2)
    tabla.style = "Table Grid"
    for etiqueta, clave in [
        ("Costo estimado", "costo_estimado"),
        ("Beneficio estimado", "beneficio_estimado"),
        ("ROI estimado", "roi_estimado"),
        ("Payback estimado", "payback_estimado"),
    ]:
        fila = tabla.add_row().cells
        fila[0].text = etiqueta
        fila[1].text = datos.get(clave) or "Por definir"
        _sombrear_celda(fila[0], _AZUL_CLARO_HEX)
        _colorear_texto_celda(fila[0], _AZUL)

    p_titulo_supuestos = doc.add_heading("Supuestos", level=1)
    for run in p_titulo_supuestos.runs:
        run.font.color.rgb = _AZUL
    _borde_inferior(p_titulo_supuestos, _AZUL_CLARO_HEX)

    supuestos = datos.get("supuestos") or []
    if supuestos:
        for s in supuestos:
            doc.add_paragraph(s, style="List Bullet")
    else:
        doc.add_paragraph("No especificados")

    doc.add_paragraph()
    p_recomendacion = doc.add_paragraph()
    p_recomendacion.add_run("Recomendación: ").bold = True
    recomendacion_texto = datos.get("recomendacion") or "Pendiente de análisis"
    run_recomendacion = p_recomendacion.add_run(recomendacion_texto)
    run_recomendacion.bold = True
    texto_lower = recomendacion_texto.lower()
    es_go = texto_lower.startswith("go") and not texto_lower.startswith("no go")
    run_recomendacion.font.color.rgb = _VERDE if es_go else _NARANJA_OSCURO

    doc.save(ruta_salida)


GENERADORES = {
    "charter": generar_charter_docx,
    "bpmn": generar_bpmn_docx,
    "onepager": generar_onepager_docx,
    "raci": generar_raci_docx,
    "bmc": generar_bmc_docx,
    "business_case": generar_business_case_docx,
}
